# -*- coding: utf-8 -*-
"""Generate functional capabilities Word document for cms-monitoring-service."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths_cm: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], "D9E2F3")
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row_data in rows:
        row = table.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = value
    if col_widths_cm:
        for row in table.rows:
            for i, width in enumerate(col_widths_cm):
                row.cells[i].width = Cm(width)
    return table


def main() -> None:
    out_path = Path(__file__).resolve().parents[1] / "Функциональные_возможности_cms-monitoring-service.docx"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

    title = doc.add_heading("Перечень функциональных возможностей", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Сервис: CMS Monitoring Service (cms-monitoring-service)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)

    date_p = doc.add_paragraph(f"Дата формирования: {date.today().strftime('%d.%m.%Y')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Назначение сервиса", level=1)
    doc.add_paragraph(
        "CMS Monitoring Service — backend-сервис платформы Trustflow для управления рисковыми объектами, "
        "конфигурациями интеграций с внешними системами, очередью результатов мониторинга и фоновой "
        "обработкой pull-интеграций. Сервис предоставляет REST API (префикс /api), публикует события в Kafka "
        "и взаимодействует с compliance-auth-service (cms-auth-service) для проверки прав доступа."
    )

    doc.add_heading("2. Роли пользователей", level=1)
    doc.add_paragraph(
        "В платформе используются роли, хранимые в compliance-auth-service и cms-company-info:"
    )
    add_table(
        doc,
        ["Роль", "Описание", "Доступ к функциям мониторинга по умолчанию"],
        [
            [
                "EXECUTIVE",
                "Руководитель / администратор компании",
                "Полный доступ: при регистрации пользователю назначаются все permissions, "
                "включая MANAGE_INTEGRATIONS, VIEW_INTEGRATIONS_PAGE, MANAGE_RISK_OBJECTS",
            ],
            [
                "SUPERVISOR",
                "Супервайзер",
                "Доступ определяется индивидуально назначенными permissions (по умолчанию — пустой набор)",
            ],
            [
                "MANAGER",
                "Менеджер",
                "Доступ определяется индивидуально назначенными permissions (по умолчанию — пустой набор)",
            ],
            [
                "DEFAULT",
                "Роль по умолчанию для новых пользователей",
                "Доступ определяется индивидуально назначенными permissions (по умолчанию — пустой набор)",
            ],
        ],
        [3.0, 4.5, 8.0],
    )

    doc.add_paragraph()
    doc.add_heading("3. Модель контроля доступа", level=1)
    doc.add_paragraph(
        "Большинство операций требуют JWT Bearer-токена (глобальный AuthenticationGuard). "
        "Для ряда операций дополнительно выполняется проверка permission через "
        "GET /api/users/{userId}/permissions/check в compliance-auth-service. "
        "Permissions назначаются на уровне пользователя (не жёстко привязаны к роли), "
        "однако EXECUTIVE при создании получает полный набор прав."
    )
    doc.add_paragraph("Используемые в сервисе permissions:", style="List Bullet")
    for perm, desc in [
        ("VIEW_INTEGRATIONS_PAGE", "просмотр списка и карточки интеграций"),
        ("MANAGE_INTEGRATIONS", "создание, изменение, удаление интеграций и смена статуса"),
        ("MANAGE_RISK_OBJECTS", "удаление рисковых объектов"),
    ]:
        p = doc.add_paragraph(style="List Bullet 2")
        p.add_run(f"{perm} — ").bold = True
        p.add_run(desc)

    doc.add_paragraph()
    doc.add_paragraph("Условные обозначения в таблицах ниже:", style="List Bullet")
    doc.add_paragraph(
        "«Все роли с permission» — MANAGER / SUPERVISOR / DEFAULT / EXECUTIVE при наличии соответствующего permission у пользователя.",
        style="List Bullet 2",
    )
    doc.add_paragraph(
        "«Любой аутентифицированный пользователь» — валидный JWT; отдельная проверка permission в коде не выполняется.",
        style="List Bullet 2",
    )
    doc.add_paragraph(
        "«Системный (межсервисный)» — публичный эндпоинт без JWT; доступ только с localhost.",
        style="List Bullet 2",
    )

    doc.add_heading("4. REST API: рисковые объекты", level=1)
    add_table(
        doc,
        ["№", "Функция", "HTTP / путь", "Требуемый permission", "Роли с доступом"],
        [
            ["4.1", "Создание рискового объекта", "POST /api/risk-objects", "—", "Любой аутентифицированный пользователь"],
            ["4.2", "Список рисковых объектов (постранично, поиск по имени)", "GET /api/risk-objects", "—", "Любой аутентифицированный пользователь"],
            ["4.3", "Краткий список моделей (id, name)", "GET /api/risk-object-models", "—", "Любой аутентифицированный пользователь"],
            ["4.4", "Детали рискового объекта по id", "GET /api/risk-objects/{id}", "—", "Любой аутентифицированный пользователь"],
            ["4.5", "Обновление рискового объекта", "PUT /api/risk-objects/{id}", "—", "Любой аутентифицированный пользователь"],
            ["4.6", "Смена статуса рискового объекта", "PUT /api/risk-objects/{id}/status", "—", "Любой аутентифицированный пользователь"],
            [
                "4.7",
                "Удаление рискового объекта",
                "DELETE /api/risk-objects/{id}",
                "MANAGE_RISK_OBJECTS",
                "EXECUTIVE (по умолчанию); MANAGER / SUPERVISOR / DEFAULT — при назначении permission",
            ],
            ["4.8", "История изменений рисковых объектов (список)", "GET /api/risk-objects/change-history", "—", "Любой аутентифицированный пользователь"],
            ["4.9", "Детали записи истории изменений", "GET /api/risk-objects/change-history/{historyId}", "—", "Любой аутентифицированный пользователь"],
            [
                "4.10",
                "Получение рискового объекта по UUID (межсервисный)",
                "GET /api/internal/risk-objects/{uuid}",
                "—",
                "Системный (межсервисный); без роли пользователя",
            ],
        ],
        [1.0, 4.5, 4.5, 3.5, 5.0],
    )

    doc.add_paragraph()
    doc.add_heading("5. REST API: конфигурации интеграций", level=1)
    add_table(
        doc,
        ["№", "Функция", "HTTP / путь", "Требуемый permission", "Роли с доступом"],
        [
            [
                "5.1",
                "Создание конфигурации интеграции (push/pull/broker)",
                "POST /api/integration-configs",
                "MANAGE_INTEGRATIONS",
                "EXECUTIVE (по умолчанию); остальные роли — при назначении permission",
            ],
            [
                "5.2",
                "Таблица интеграций (постранично, поиск по имени)",
                "GET /api/integration-configs",
                "VIEW_INTEGRATIONS_PAGE",
                "EXECUTIVE (по умолчанию); остальные роли — при назначении permission",
            ],
            [
                "5.3",
                "Просмотр интеграции по id",
                "GET /api/integration-configs/{id}",
                "VIEW_INTEGRATIONS_PAGE",
                "EXECUTIVE (по умолчанию); остальные роли — при назначении permission",
            ],
            [
                "5.4",
                "Сохранение изменений интеграции",
                "PUT /api/integration-configs/{id}",
                "MANAGE_INTEGRATIONS",
                "EXECUTIVE (по умолчанию); остальные роли — при назначении permission",
            ],
            [
                "5.5",
                "Удаление интеграции",
                "DELETE /api/integration-configs/{id}",
                "MANAGE_INTEGRATIONS",
                "EXECUTIVE (по умолчанию); остальные роли — при назначении permission",
            ],
            [
                "5.6",
                "Изменение флага активности / статуса интеграции",
                "PUT /api/integration-configs/{id}/status",
                "MANAGE_INTEGRATIONS",
                "EXECUTIVE (по умолчанию); остальные роли — при назначении permission",
            ],
            [
                "5.7",
                "История изменений интеграций (список, поиск)",
                "GET /api/integration-configs/change-history",
                "—",
                "Любой аутентифицированный пользователь",
            ],
            [
                "5.8",
                "Получение конфигурации интеграции по id (межсервисный)",
                "GET /api/internal/integration-configs/{id}",
                "—",
                "Системный (межсервисный); без роли пользователя",
            ],
        ],
        [1.0, 4.5, 4.5, 3.5, 5.0],
    )

    doc.add_paragraph()
    doc.add_heading("6. REST API: результаты мониторинга", level=1)
    add_table(
        doc,
        ["№", "Функция", "HTTP / путь", "Требуемый permission", "Роли с доступом"],
        [
            [
                "6.1",
                "Статистика очереди monitoring_result и monitoring_retry",
                "GET /api/monitoring-results/statistics",
                "—",
                "Любой аутентифицированный пользователь",
            ],
            [
                "6.2",
                "Взять запись мониторинга в обработку (удаление из очереди)",
                "PUT /api/monitoring-results/{id}/take",
                "—",
                "Системный вызов (без JWT); для фоновых воркеров платформы",
            ],
        ],
        [1.0, 4.5, 4.5, 3.5, 5.0],
    )

    doc.add_paragraph()
    doc.add_heading("7. Фоновые и системные возможности (без UI-роли)", level=1)
    add_table(
        doc,
        ["№", "Функция", "Описание", "Инициатор"],
        [
            [
                "7.1",
                "Периодический опрос pull-интеграций",
                "POST на endpointUrl внешних систем по расписанию",
                "Сервис (фоновый процесс)",
            ],
            [
                "7.2",
                "Обработка outbox-очереди входящих данных",
                "Сохранение и маршрутизация данных мониторинга",
                "Сервис (фоновый процесс)",
            ],
            [
                "7.3",
                "Публикация событий в Kafka (ws_events, risk_topic)",
                "Статусы интеграций, счётчики вызовов, пайплайн рисков",
                "Сервис (фоновый процесс)",
            ],
            [
                "7.4",
                "Миграции схемы PostgreSQL",
                "Автоматический запуск при старте сервиса",
                "Сервис при деплое",
            ],
            [
                "7.5",
                "Swagger-документация API",
                "GET /api/docs",
                "Разработчик / тестировщик (без проверки роли в коде)",
            ],
        ],
        [1.0, 4.0, 7.5, 4.0],
    )

    doc.add_paragraph()
    doc.add_heading("8. Сводная матрица: permission → роли", level=1)
    add_table(
        doc,
        ["Permission", "EXECUTIVE", "SUPERVISOR", "MANAGER", "DEFAULT"],
        [
            ["VIEW_INTEGRATIONS_PAGE", "Да (по умолчанию)", "При назначении", "При назначении", "При назначении"],
            ["MANAGE_INTEGRATIONS", "Да (по умолчанию)", "При назначении", "При назначении", "При назначении"],
            ["MANAGE_RISK_OBJECTS", "Да (по умолчанию)", "При назначении", "При назначении", "При назначении"],
        ],
        [4.5, 3.0, 3.0, 3.0, 3.0],
    )

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("Примечание. ").bold = True
    note.add_run(
        "Назначение и изменение permissions для пользователей ролей MANAGER, SUPERVISOR и DEFAULT "
        "выполняется в compliance-auth-service пользователем с правом EDIT_USERS (как правило, EXECUTIVE). "
        "Источник данных: исходный код cms-monitoring-service и compliance-auth-service."
    )

    doc.save(out_path)
    print(f"Created: {out_path}")


if __name__ == "__main__":
    main()
